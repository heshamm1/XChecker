from tkinter import *
from tkinter import ttk
from tkinter import filedialog
import requests
from docx import Document
import re
import base64
import urllib.parse

root = Tk()
root.title("XCh3ck3r")
root.geometry("720x400")
root.configure(bg="#17202A")

style = ttk.Style()
style.theme_use("alt")
style.configure("TFrame", background="#17202A")
style.configure("TLabel", background="#17202A", foreground="white")
style.configure("TButton", background="#17202A", foreground="Green", type="italic")

url = StringVar()
label = ttk.Label(root, text="Enter URL to check: ", foreground="White", font="bold")
entry = ttk.Entry(root, width=40, textvariable=url)

summary_label = ttk.Label(root, text="", foreground="#F4D03F", font=("bold", 11))



def get_reports():
    url_input = url.get()

    try:
        # Ensure that the input is a valid URL
        parsed_url = urllib.parse.urlparse(url_input)
        if not (parsed_url.scheme and parsed_url.netloc):
            raise ValueError("Invalid URL")

        # Fetch reports from each site
        norton_url = f"https://safeweb.norton.com/report/show?url={url_input}"
        sucuri_url = f"https://sitecheck.sucuri.net/results/{url_input}"
        xforce_url = f"https://exchange.xforce.ibmcloud.com/url/{url_input}"    
        cloudflare_url = f"https://radar.cloudflare.com/scan/{url_input}"
        talos_url = f"https://talosintelligence.com/reputation_center/lookup?search={url_input}"
        IPQualityScore_url = re.sub(r'://', '',f"www.ipqualityscore.com/threat-feeds/malicious-url-scanner/http%3A%2F%2F{url_input}")
        virustotal_url = f"https://www.virustotal.com/vtapi/v2/url/report?apikey=a4ac0da5ccd96d4c8673b95962f45ae271de61301e78bf2b5d3ade9e6ee6b866&resource={url_input}"
        urlvoid_url = re.sub(r'https?://', '', f"https://www.urlvoid.com/scan/{url_input}/")

        reports = {"Norton": norton_url,
                   "Sucuri": sucuri_url,
                   "IBM X-Force Exchange": xforce_url,
                   "Cloudflare": cloudflare_url,
                   "Talos Intelligence": talos_url,
                   "IPQualityScore": IPQualityScore_url,
                   "VirusTotal": virustotal_url,
                   "URLVoid": urlvoid_url}

        generate_report(reports)
        generate_summary(reports)
    except ValueError:
        summary_label["text"] = "Invalid URL"
    except requests.exceptions.RequestException:
        summary_label["text"] = "Error fetching reports"
        
        

def generate_report(reports):
    document = Document()
    document.add_heading("BR's URL Analysis Report", level=0)

    for site, url in reports.items():
        document.add_heading(site, level=1)
        document.add_paragraph(url)

    report_file_path = filedialog.asksaveasfilename(defaultextension=".docx")
    document.save(report_file_path)

    label["text"] = "Report generated!"

def generate_summary(reports):
    summary = {}
    for site, response in reports.items():
        if "norton" in site.lower():
            NORTON_response = base64.b64encode(response.encode('utf-8')).decode('utf-8')
            if "CAUTION" or "WARNING" in NORTON_response.lower():
                summary[site] = "NOT SAFE"
            else:
                summary[site] = "SAFE"
                
        elif "sucuri" in site.lower():
            SUCURI_response = base64.b64encode(response.encode('utf-8')).decode('utf-8')
            if "Not Detected" in SUCURI_response.lower():
                summary[site] = "NOT DETECTED"
            else: 
                summary[site] = "DETECTED"
                
        elif "ibm x-force exchange" in site.lower():
            IBM_response = base64.b64encode(response.encode('utf-8')).decode('utf-8')
            if "unknown" in IBM_response.lower():
                summary[site] = "UNKNOWN, CHECK THE REPORT"
            else:
                summary[site] = "UNKNOWN, CHECK THE REPORT"
                
        elif "talos intelligence" in site.lower():
            TI_response = base64.b64encode(response.encode('utf-8')).decode('utf-8')
            if "UnTrusted" in TI_response.lower():
                summary[site] = "UNTRUSTED"
            elif "Neural" in TI_response.lower():
                summary[site] = "NEURAL"
            else:
                summary[site] = "TRUSTED"
                
        elif "IPQualityScore" in site.lower():
            IPQS_response = base64.b64encode(response.encode('utf-8')).decode('utf-8')
            if "Not Safe" in IPQS_response.lower():
                summary[site] = "NOT SAFE"
            elif "SAFE" in response.lower():
                summary[site] = "SAFE"
            else:
                summary[site] = "UNKNOWN, CHECK THE REPORT"
                
        elif "virus total" in site.lower():
            VT_response = base64.b64encode(response.encode('utf-8')).decode('utf-8')
            if "Suspicious" in VT_response.lower():
                summary[site] = "NOT SAFE"
            else:
                summary[site] = "KINDLY CHECK THE REPORT"
                
        elif "urlvoid" in site.lower():
            URLV_response = base64.b64encode(response.encode('utf-8')).decode('utf-8')
            if "Detected" in URLV_response.lower():
                summary[site] = "NOT SAFE"
            else:
                summary[site] = "SAFE"
                
        else:
            summary[site] = "UNKNOWN, CHECK THE REPORT"
            
    summary_text = "\n".join([f"{site}=> {status}" for site, status in summary.items()])
    summary_label["text"] = summary_text    

label.grid(column=0, row=0, padx=10, pady=10)
entry.grid(column=1, row=0, padx=10, pady=10)
button = ttk.Button(root, text="Check URL", command=get_reports)
button.grid(column=2, row=0, padx=10, pady=10)

summary_label.grid(column=0, row=1, columnspan=3, padx=10, pady=10)

made_by_label = ttk.Label(root, text="Made by: Hesham Ahmed", foreground="#ABB2B9" )
made_by_label.grid(column=0, row=3, columnspan=4, padx=11, pady=11)
team_label = ttk.Label(root, text="Business Resilience Team", foreground="#ABB2B9")
team_label.grid(column=0, row=2, columnspan=3, padx=11, pady=11 )
fire_lable = ttk.Label(root, text="FIRE IT, BROTHER!", foreground="Red")
fire_lable.grid(column=0, row=5, columnspan=3, padx=11, pady=11 )
root.resizable(False,False)
root.mainloop()
